# -*- coding: utf-8 -*-
import csv
import re
from pathlib import Path
import tkinter as tk
from tkinter import filedialog, messagebox
import argparse
import sys
# O 'os' foi removido, pois não há mais manipulação de arquivos temporários

# O suporte a DOC e pywin32 foi removido, deixando apenas a MockDocument como fallback.
class MockDocument:
    def __init__(self, path=None): pass
    def save(self, path): pass
    
# Adicionado para suporte a arquivos DOCX (requer a biblioteca python-docx instalada)
try:
    from docx import Document
    from docx.shared import Inches
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run
    # Adicionando a importação do docx.oxml para anexar conteúdo (complexo, mas necessário)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from copy import deepcopy # Necessário para a cópia segura de elementos XML
except ImportError:
    # Esta exceção é esperada caso o python-docx não esteja instalado.
    Document = MockDocument


# --- Funções de Leitura de Arquivos com Fallback de Codificação ---
def ler_arquivo_com_fallback(caminho):
    if not Path(caminho).exists():
        raise FileNotFoundError(f"O arquivo não foi encontrado no caminho: {caminho}")

    for enc in ['utf-8-sig', 'utf-8', 'windows-1252', 'latin-1']:
        try:
            return Path(caminho).read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    raise UnicodeDecodeError(f"Não foi possível ler o arquivo {caminho} com codificações conhecidas.")

# --- Funções de Carregamento de Dados (MÚLTIPLOS CONJUNTOS) ---

def carregar_todos_os_conjuntos(caminho_csv):
    """
    Carrega múltiplos conjuntos de dados (Dados Gerais e Pontos) de um único 
    arquivo CSV, separados por '**'.
    
    Retorna uma lista de tuplas: [(dados_gerais_1, pontos_1), (dados_gerais_2, pontos_2), ...]
    """
    todos_conjuntos = []
    
    try:
        print(f"DEBUG: Carregando múltiplos conjuntos do CSV ÚNICO: {caminho_csv}")
        conteudo = ler_arquivo_com_fallback(caminho_csv)
        
        # Divide o conteúdo do CSV em blocos usando o separador '**'
        blocos_de_dados = re.split(r'\s*\*\*\s*', conteudo) # Splita por ** e remove espaços/quebras de linha ao redor
        
        for idx_bloco, bloco in enumerate(blocos_de_dados):
            bloco_linhas = [l.strip() for l in bloco.split('\n') if l.strip()]
            if not bloco_linhas:
                continue

            dados_gerais = {}
            pontos = []
            cabecalho = []
            lendo_pontos = False
            
            print(f"DEBUG: Processando bloco {idx_bloco + 1}...")

            # Processa as linhas dentro de cada bloco
            for i, linha in enumerate(bloco_linhas):
                # Tenta usar o leitor CSV para garantir o delimitador e aspas
                try:
                    # Tenta ler com delimitador ';'. Assume o formato do arquivo original.
                    campos = next(csv.reader([linha], delimiter=';')) 
                    campos = [campo.strip().strip('"') for campo in campos]
                    # Limpa campos vazios no final
                    while campos and not campos[-1]:
                        campos.pop()
                except Exception:
                    # Fallback para split simples
                    campos = [campo.strip().strip('"') for campo in linha.split(';')]

                if not campos or not campos[0]:
                    continue

                # ----------------------------------------------------
                # 1. Identificação/Início da seção de PONTOS
                # ----------------------------------------------------
                if campos[0] == "<PONTO>":
                    if lendo_pontos: # Ignora cabeçalhos <PONTO> repetidos dentro do mesmo bloco
                        continue 
                    lendo_pontos = True
                    cabecalho = campos
                    continue

                # ----------------------------------------------------
                # 2. Leitura de PONTOS
                # ----------------------------------------------------
                if lendo_pontos:
                    try:
                        num_campos_esperados = len(cabecalho)
                        if num_campos_esperados == 0:
                            raise ValueError(f"Cabeçalho '<PONTO>' está vazio ou malformado no bloco {idx_bloco + 1}.")

                        if len(campos) < num_campos_esperados:
                            campos.extend([''] * (num_campos_esperados - len(campos)))
                        elif len(campos) > num_campos_esperados:
                            campos = campos[:num_campos_esperados]

                        ponto = dict(zip(cabecalho, [val.strip() for val in campos]))
                        pontos.append(ponto)
                    except Exception as e:
                        print(f"AVISO: Linha ignorada na seção de pontos do bloco {idx_bloco + 1} por erro de formato: {e}")
                        
                # ----------------------------------------------------
                # 3. Leitura de DADOS GERAIS
                # ----------------------------------------------------
                else: 
                    try:
                        # O formato esperado para dados gerais é <CHAVE>;<VALOR>
                        if len(campos) >= 2 and re.match(r"<[A-Z_]+>", campos[0]):
                            chave = campos[0]
                            valor = campos[1]
                            dados_gerais[chave] = valor
                    except IndexError:
                        continue # Linha ignorada se não for um par CHAVE;VALOR válido

            # --- Validação e Finalização do Bloco ---
            # Apenas adiciona se houver dados gerais OU pontos.
            if not dados_gerais and not pontos:
                continue
                
            if len(pontos) < 2:
                print(f"AVISO: Bloco {idx_bloco + 1} sem pontos de perímetro (ou menos de 2), mas com dados gerais. Será processado como arquivo simples.")
                
            # Lógica para evitar ponto final duplicado (se o último ponto for igual ao primeiro)
            if len(pontos) > 1 and pontos[0].get('<PONTO>') == pontos[-1].get('<PONTO>'):
                pontos.pop()
                
            # Lógica de fallback para Área (CHAVE CORRIGIDA)
            if "<AREAHE>" in dados_gerais and "<AREAM2>" not in dados_gerais:
                dados_gerais["<AREAM2>"] = dados_gerais["<AREAHE>"] 
                
            # Adiciona o conjunto à lista
            todos_conjuntos.append((dados_gerais, pontos))
        
        if not todos_conjuntos:
            raise ValueError("Nenhum conjunto de dados (Gerais + Pontos) válido foi encontrado no CSV único.")
            
        print(f"DEBUG: Carregamento concluído. {len(todos_conjuntos)} conjunto(s) de dados pronto(s) para processamento.")
        return todos_conjuntos

    except Exception as e:
        raise ValueError(f"Erro fatal ao carregar dados do CSV ÚNICO: {e}")

# --- Funções Auxiliares (Lógica de Substituição de Texto) ---

def normalizar_chaves_rtf(texto_rtf, chaves):
    """
    Normaliza as chaves no texto RTF, removendo códigos de formatação 
    quebras de linha/espaços que o editor pode ter inserido dentro da chave.
    """
    for chave in chaves:
        padrao = ''
        for letra in chave:
            # Captura a letra, seguida por zero ou mais sequências de códigos RTF (\pard, \b0, etc.), 
            # espaços, ou chaves de agrupamento (quebra o token)
            padrao += re.escape(letra) + r'(?:\\[a-z]+\d* ?|[\s{}])*?'
        padrao_regex = re.compile(padrao, flags=re.IGNORECASE)
        texto_rtf = padrao_regex.sub(chave, texto_rtf)
    return texto_rtf

def replace_placeholder_in_paragraph(paragraph, placeholder, replacement):
    """
    Substitui um placeholder em um parágrafo do DOCX, mantendo a formatação
    original do texto que o envolve.
    """
    if placeholder not in paragraph.text:
        return

    text_before = ''
    start_run_index = -1
    start_offset = -1
    current_char_count = 0
    for i, run in enumerate(paragraph.runs):
        run_text = run.text
        run_len = len(run.text)
        if current_char_count + run_len > paragraph.text.find(placeholder) and start_run_index == -1:
            start_run_index = i
            start_pos_in_para = paragraph.text.find(placeholder)
            start_offset = start_pos_in_para - current_char_count
            break
        current_char_count += run_len
    if start_run_index == -1: return

    end_run_index = start_run_index
    end_offset = -1
    char_count = current_char_count
    placeholder_end_pos = paragraph.text.find(placeholder) + len(placeholder)
    for i in range(start_run_index, len(paragraph.runs)):
        run = paragraph.runs[i]
        run_len = len(run.text)
        if char_count + run_len >= placeholder_end_pos:
            end_run_index = i
            end_offset = placeholder_end_pos - char_count
            break
        char_count += run_len
    if end_run_index == -1: return
        
    start_run = paragraph.runs[start_run_index]
    end_run = paragraph.runs[end_run_index]
    
    prefix_text = start_run.text[:start_offset]
    suffix_text = end_run.text[end_offset:]
    
    # 1. Caso a chave esteja em uma única run
    if start_run_index == end_run_index:
        start_run.text = prefix_text + str(replacement) + suffix_text
    else:
        # 2. Caso a chave esteja distribuída por múltiplas runs
        # Substitui o início da run inicial
        start_run.text = prefix_text + str(replacement)
        
        # Limpa as runs intermediárias
        for i in range(start_run_index + 1, end_run_index):
            paragraph.runs[i].clear()
            
        # Substitui o final da run final
        end_run.text = suffix_text
        
        # Limpa runs vazias que sobraram
        for i in range(len(paragraph.runs) - 1, -1, -1):
            if not paragraph.runs[i].text:
                p = paragraph._element
                r = paragraph.runs[i]._element
                p.remove(r)
                
def aplicar_estilo_base(run_alvo, run_origem):
    """Copia fonte, tamanho e negrito de uma run para outra."""
    run_alvo.font.name = run_origem.font.name if run_origem.font.name else "Arial"
    if run_origem.font.size:
        run_alvo.font.size = run_origem.font.size
    run_alvo.bold = run_origem.bold
    
def replace_placeholder_docx(paragrafo, dados):
    """Substitui placeholders mantendo a formatação de cada run individualmente."""
    for chave, valor in dados.items():
        if chave in paragrafo.text:
            for run in paragrafo.runs:
                if chave in run.text:
                    run.text = run.text.replace(chave, str(valor))
                    
def substituir_texto_em_docx(estrutura, dados_substituicao):
    """
    Aplica as substituições em todos os parágrafos e tabelas da estrutura fornecida
    (Documento ou Célula de Tabela).
    """
    def substituir_em_estrutura_recursiva(estrutura):
        for paragrafo in estrutura.paragraphs:
            for chave, valor in dados_substituicao.items():
                replace_placeholder_in_paragraph(paragrafo, chave, str(valor))
    
    substituir_em_estrutura_recursiva(estrutura)

    if hasattr(estrutura, 'tables'):
        for tabela in estrutura.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    substituir_em_estrutura_recursiva(celula)
                
def criar_padrao_rtf_flexivel(texto_alvo):
    """
    Cria um padrão regex flexível que ignora códigos RTF, usado para 
    localizar o trecho de "confrontante repetido" no bloco de repetição.
    """
    padrao = ''
    for letra in texto_alvo:
        if letra.isspace():
            padrao += r'\s*' 
        else:
            # Captura a letra, seguida por zero ou mais sequências de códigos RTF ou chaves
            padrao += re.escape(letra) + r'(?:\\[a-z]+\d* ?|[\s{}])*?'
    return padrao

# --- Funções de Processamento de Conteúdo de Bloco (RTF e DOCX) ---

def processar_rtf_string_content(modelo_rtf, dados_gerais, pontos, ignorar_confrontante_repetido=False):
    """
    Processa um ÚNICO memorial RTF e retorna a STRING do conteúdo, 
    garantindo que os nomes dos pontos (<PONTO>) fiquem em negrito.
    """
    try:
        # Tenta ler o RTF com uma codificação segura
        texto_modelo = Path(modelo_rtf).read_text(encoding='windows-1252', errors='ignore')

        if not texto_modelo:
            raise ValueError("O arquivo modelo RTF está vazio ou não pôde ser lido.")
        
        # 0. Normalização das chaves do template
        chaves_a_normalizar = list(dados_gerais.keys()) + list(pontos[0].keys() if pontos else []) + ['<***>', '<CONFRO>']
        # Nota: Assume-se que a função normalizar_chaves_rtf existe no escopo global
        texto_pronto = normalizar_chaves_rtf(texto_modelo, chaves_a_normalizar)
        
        # 1. Substitui dados gerais (em todo o documento)
        for chave, valor in dados_gerais.items():
            texto_pronto = texto_pronto.replace(chave, str(valor))

        # 2. Encontra o bloco de repetição
        padrao_bloco = r"(?s)<\*\*\*>(.*?)<\*\*\*>"
        match_bloco = re.search(padrao_bloco, texto_pronto)

        if not match_bloco or len(pontos) < 2:
            texto_pronto = texto_pronto.replace('<***>', '')
            # Aplica negrito no PONTO inicial mesmo se não houver bloco
            if pontos:
                ponto_nome = pontos[0].get('<PONTO>', '')
                texto_pronto = texto_pronto.replace('<PONTO>', f"{{\\b {ponto_nome}}}")
            return texto_pronto

        # Separar o texto em partes
        texto_antes = texto_pronto[:match_bloco.start()]
        bloco_base = match_bloco.group(1).replace('<***>', '')
        texto_depois = texto_pronto[match_bloco.end():]
        blocos_gerados = []
        
        # --- ABERTURA: ponto inicial (M01) ---
        first_point = pontos[0]
        for key, value in first_point.items():
            val_to_replace = value
            if key == '<PONTO>':
                # Formatação RTF para Negrito: {\b texto}
                val_to_replace = f"{{\\b {value}}}"
            texto_antes = texto_antes.replace(key, str(val_to_replace))

        # --- GERAÇÃO DOS PARÁGRAFOS DE TRANSIÇÃO (BLOCO DE REPETIÇÃO) ---
        last_confrontante = None
        # Nota: Assume-se que criar_padrao_rtf_flexivel existe no escopo global
        frase_alvo_rtf = " confrontando com a propriedade de <CONFRONTANTE>"
        padrao_remover_frase = criar_padrao_rtf_flexivel(frase_alvo_rtf) + r'\s*[,\.]*'

        for idx in range(len(pontos) - 1):
            current_point = pontos[idx]
            next_point = pontos[idx + 1]
            confronto_num = idx + 1

            formatted_block = bloco_base[:]
            paragraph_data = current_point.copy()
            
            # Formata o ponto de destino em negrito para o RTF
            ponto_destino_nome = next_point.get('<PONTO>', '')
            ponto_destino_negrito = f"{{\\b {ponto_destino_nome}}}"

            paragraph_data.update({
                '<PONTO>': ponto_destino_negrito,
                '<UTMX>': next_point.get('<UTMY>', ''), # Mantendo lógica de inversão do original
                '<UTMY>': next_point.get('<UTMX>', ''),
            })
            
            current_confrontante_name = current_point.get('<CONFRONTANTE>', '').strip()
            
            # Tratamento de Confrontante Repetido
            if ignorar_confrontante_repetido: 
                if current_confrontante_name and current_confrontante_name == last_confrontante:
                    formatted_block, subs = re.subn(padrao_remover_frase, "", formatted_block, flags=re.IGNORECASE | re.DOTALL)
                    if subs > 0:
                        paragraph_data['<CONFRONTANTE>'] = ""
                    else:
                        paragraph_data['<CONFRONTANTE>'] = "o mesmo"
                else:
                    paragraph_data['<CONFRONTANTE>'] = current_confrontante_name
                last_confrontante = current_confrontante_name
            else: 
                paragraph_data['<CONFRONTANTE>'] = current_confrontante_name

            # Realiza as substituições no bloco gerado
            for key, value in paragraph_data.items():
                formatted_block = formatted_block.replace(key, str(value))

            if '<CONFRO>' in formatted_block:
                formatted_block = formatted_block.replace("<CONFRO>", f"Confronto {confronto_num}")

            blocos_gerados.append(formatted_block)

        # --- FECHAMENTO: (último → primeiro) ---
        last_point_of_perimeter = pontos[-1]
        first_point = pontos[0]
        confronto_num_fechamento = len(pontos) 
        
        texto_depois_final = texto_depois[:]
        fechamento_data = last_point_of_perimeter.copy()

        # Ponto de destino final (volta ao início) em negrito
        ponto_inicio_nome = first_point.get('<PONTO>', '')
        ponto_inicio_negrito = f"{{\\b {ponto_inicio_nome}}}"

        fechamento_data['<PONTO>'] = ponto_inicio_negrito
        fechamento_data['<UTMX>'] = first_point.get('<UTMX>', '')
        fechamento_data['<UTMY>'] = first_point.get('<UTMY>', '')

        # Aplica substituições no fechamento
        for key, value in fechamento_data.items():
            texto_depois_final = texto_depois_final.replace(key, str(value))

        if '<CONFRO>' in texto_depois_final:
            texto_depois_final = texto_depois_final.replace('<CONFRO>', f"Confronto {confronto_num_fechamento}")
            
        # Monta o texto final
        texto_final = texto_antes + "".join(blocos_gerados) + texto_depois_final
        return texto_final

    except Exception as e:
        raise Exception(f"Erro ao processar o bloco RTF. Detalhes: {e}")
    
def processar_docx_bloco(caminho, dados_gerais, pontos, ignorar_rep=False):
    doc = Document(caminho)
    substituir_texto_em_docx(doc, dados_gerais)

    start_idx = -1
    end_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if "<***>" in p.text:
            if start_idx == -1: start_idx = i
            end_idx = i

    if start_idx == -1 or len(pontos) < 2:
        for p in doc.paragraphs: replace_placeholder_docx(p, {"<***>": ""})
        return doc

    # --- IDENTIFICAÇÃO DOS TEMPLATES ---
    p_intro = doc.paragraphs[start_idx]
    p_outro = doc.paragraphs[end_idx]
    
    font_name = p_intro.runs[0].font.name if p_intro.runs else "Arial"
    font_size = p_intro.runs[0].font.size if p_intro.runs else Pt(11)

    # Captura o texto do bloco de repetição
    bloco_raw = ""
    for i in range(start_idx, end_idx + 1):
        bloco_raw += doc.paragraphs[i].text + "\n"
    match = re.search(r"(?s)<\*\*\*>(.*?)<\*\*\*>", bloco_raw)
    bloco_template_txt = match.group(1).strip() if match else ""

    # --- 1. PROCESSAMENTO DA ABERTURA (Especial para P1 e P2 em negrito) ---
    
    val_p1 = str(pontos[0].get('<PONTO>', ''))
    val_n1 = str(pontos[0].get('<UTMY>', ''))
    val_e1 = str(pontos[0].get('<UTMX>', ''))
    
    val_p2 = str(pontos[1].get('<PONTO>', ''))
    val_n2 = str(pontos[1].get('<UTMY>', ''))
    val_e2 = str(pontos[1].get('<UTMX>', ''))

    # Divide o parágrafo de introdução pelo marcador de repetição
    partes_intro_raw = p_intro.text.split("<***>")
    part_a_txt = partes_intro_raw[0]
    part_b_txt = partes_intro_raw[1] if len(partes_intro_raw) > 1 else ""

    # Substitui dados na Parte A (P1)
    for k, v in pontos[0].items(): part_a_txt = part_a_txt.replace(k, str(v))
    part_a_txt = part_a_txt.replace("< PONTO>", val_p1)
    part_a_txt = re.sub(r'<\s*UTMY\s*>', val_n1, part_a_txt)
    part_a_txt = re.sub(r'<\s*UTMX\s*>', val_e1, part_a_txt)

    # Substitui dados na Parte B (P2)
    dados_p2 = pontos[0].copy()
    dados_p2.update({'<PONTO>': val_p2, '<UTMY>': val_n2, '<UTMX>': val_e2})
    for k, v in dados_p2.items(): part_b_txt = part_b_txt.replace(k, str(v))
    part_b_txt = part_b_txt.replace("< PONTO>", val_p2)
    part_b_txt = re.sub(r'<\s*UTMY\s*>', val_n2, part_b_txt)
    part_b_txt = re.sub(r'<\s*UTMX\s*>', val_e2, part_b_txt)

    # Limpa e reconstrói o parágrafo com estilos
    p_intro.clear()

    def add_styled_runs(para, text, p_val, n_val, e_val):
        # Regex para capturar os valores que devem ser negrito
        # Inclui o nome do ponto e as coordenadas com prefixos N e E
        regex_bold = f"({re.escape(p_val)}|N\s*=?\s*{re.escape(n_val)}|E\s*=?\s*{re.escape(e_val)})"
        segs = re.split(regex_bold, text)
        for s in segs:
            if not s: continue
            r = para.add_run(s)
            r.font.name = font_name
            r.font.size = font_size
            # Verifica se o segmento atual é um dos valores de destaque
            if s == p_val or (s.startswith('N') and n_val in s) or (s.startswith('E') and e_val in s):
                r.bold = True

    add_styled_runs(p_intro, part_a_txt, val_p1, val_n1, val_e1)
    if part_b_txt:
        add_styled_runs(p_intro, part_b_txt, val_p2, val_n2, val_e2)

    # --- 2. PROCESSAMENTO DAS REPETIÇÕES (PONTOS INTERMEDIÁRIOS) ---
    ancora = p_outro
    ultimo_conf = None
    if len(pontos) > 2:
        for i in range(1, len(pontos) - 1):
            p_at = pontos[i]
            p_pr = pontos[i+1]
            txt_bloco = bloco_template_txt[:]
            
            v_p = str(p_pr.get('<PONTO>', ''))
            v_n = str(p_pr.get('<UTMY>', ''))
            v_e = str(p_pr.get('<UTMX>', ''))

            # Substitui chaves comuns
            for k, v in p_at.items():
                if k not in ['<PONTO>', '<UTMX>', '<UTMY>']:
                    txt_bloco = txt_bloco.replace(k, str(v))
            txt_bloco = txt_bloco.replace("<CONFRO>", f"Confronto {i+1}")

            # Identifica tokens de negrito no template
            txt_bloco = re.sub(r'<\s*PONTO\s*>', '##B_P##', txt_bloco, flags=re.I)
            txt_bloco = re.sub(r'<\s*UTMY\s*>', '##B_N##', txt_bloco, flags=re.I)
            txt_bloco = re.sub(r'<\s*UTMX\s*>', '##B_E##', txt_bloco, flags=re.I)

            p_n = ancora.insert_paragraph_before("")
            p_n.style = p_intro.style
            
            partes = re.split(r'([NE]\s*=?\s*##B_[NE]##|##B_P##)', txt_bloco)
            for pt in partes:
                if not pt: continue
                r = p_n.add_run()
                r.font.name = font_name
                r.font.size = font_size
                if "##B_P##" in pt:
                    r.text = pt.replace("##B_P##", v_p); r.bold = True
                elif "##B_N##" in pt:
                    r.text = pt.replace("##B_N##", v_n); r.bold = True
                elif "##B_E##" in pt:
                    r.text = pt.replace("##B_E##", v_e); r.bold = True
                else:
                    r.text = pt

    # --- 3. PROCESSAMENTO DO FECHAMENTO ---
    for run in p_outro.runs:
        if "<***>" in run.text: run.text = run.text.replace("<***>", "")
    
    v_p_f = str(pontos[0].get('<PONTO>', ''))
    v_n_f = str(pontos[0].get('<UTMY>', ''))
    v_e_f = str(pontos[0].get('<UTMX>', ''))

    d_f = pontos[-1].copy()
    d_f.update({'<PONTO>': v_p_f, '<UTMX>': v_e_f, '<UTMY>': v_n_f, '<CONFRO>': f"Confronto {len(pontos)}"})
    replace_placeholder_docx(p_outro, d_f)
    
    # Aplica negrito manual no fechamento para garantir P1
    for run in p_outro.runs:
        if v_p_f in run.text or v_n_f in run.text or v_e_f in run.text:
            run.bold = True

    if end_idx > start_idx + 1:
        for i in range(end_idx - 1, start_idx, -1):
            p_m = doc.paragraphs[i]
            p_m._element.getparent().remove(p_m._element)

    final_subs = dados_gerais.copy()
    final_subs.update(pontos[0])
    substituir_texto_em_docx(doc, final_subs)
    
    return doc
    
    
def processar_modelo(caminho_modelo, todos_conjuntos, saida_caminho, ignorar_confrontante_repetido=False):
    """
    Processa o modelo com MÚLTIPLOS conjuntos de dados e salva em um único arquivo de saída.
    CORREÇÃO RTF: Implementa uma lógica mais robusta para extrair e combinar o corpo do RTF, 
    garantindo que o cabeçalho/rodapé seja apenas do primeiro bloco.
    """
    caminho_modelo_path = Path(caminho_modelo)
    extensao = caminho_modelo_path.suffix.lower()
    
    if not todos_conjuntos:
        print("Aviso: Nenhum conjunto de dados para processar.")
        return

    # --- Lógica RTF (Concatenação de Strings) ---
    if extensao == '.rtf':
        print(f"\n--- Iniciando Geração do RTF de Múltiplos Memoriais ({len(todos_conjuntos)} total) ---")
        
        def extrair_corpo_rtf(texto_rtf):
            """
            Remove a chave RTF de abertura ({\rtf1...}) e a chave de fechamento final (})
            do bloco RTF gerado.
            """
            # 1. Remove a chave de fechamento final '}' e espaços/novas linhas
            texto_sem_fim = re.sub(r'\}\s*$', '', texto_rtf.strip())
            
            # 2. Remove o cabeçalho RTF de abertura '{\rtf1...' (procurando pelo \rtf1 até um espaço ou {)
            # O .*? garante que pegamos o mínimo necessário até o primeiro \s ou { após o \rtf1
            match_inicio = re.match(r'\{\\rtf1.*?\s*', texto_sem_fim, re.DOTALL | re.IGNORECASE)
            
            if match_inicio:
                corpo = texto_sem_fim[match_inicio.end():]
                # Remove o que possa ter sobrado da chave de abertura após o \rtf1
                return corpo.lstrip('{').lstrip()
            else:
                # Fallback, retorna o texto sem a última chave
                return texto_sem_fim

        
        # --- 1. Processa o PRIMEIRO conjunto (retorna o RTF completo) ---
        dados_gerais_0, pontos_0 = todos_conjuntos[0]
        try:
            # O primeiro bloco é gerado com o RTF completo, incluindo o cabeçalho/rodapé
            texto_base_processado = processar_rtf_string_content(caminho_modelo, dados_gerais_0, pontos_0, ignorar_confrontante_repetido)
        except Exception as e:
            raise Exception(f"Falha fatal ao processar o Bloco 1 (RTF): {e}")

        
        # 2. Separa o cabeçalho/corpo do rodapé (chave de fechamento) do PRIMEIRO bloco
        match_end = re.search(r'\}\s*$', texto_base_processado.strip())
        if match_end:
            texto_final_corpo = texto_base_processado[:match_end.start()]
            rodape_rtf = match_end.group(0).strip() # Deve ser '}'
        else:
            texto_final_corpo = texto_base_processado
            rodape_rtf = '}'
            print("AVISO: Chave de fechamento RTF ('}') não encontrada. Tentando adicionar '}'.")


        # --- 3. Processa os conjuntos subsequentes (índice 1 em diante) ---
        if len(todos_conjuntos) > 1:
            for idx in range(1, len(todos_conjuntos)):
                item_conjunto = todos_conjuntos[idx]
                
                if not (isinstance(item_conjunto, (list, tuple)) and len(item_conjunto) == 2):
                    print(f"ERRO GRAVE DE DADOS (BLOCO {idx + 1}): Estrutura corrompida. O processo continuará.")
                    continue
                
                dados_gerais, pontos = item_conjunto

                try:
                    print(f"DEBUG: Anexando Bloco RTF {idx + 1}...")
                    
                    # Gera o RTF completo para o bloco
                    texto_memorial_completo = processar_rtf_string_content(caminho_modelo, dados_gerais, pontos, ignorar_confrontante_repetido)

                    # Extrai apenas o corpo (remove {\rtf1...} e })
                    corpo_bloco = extrair_corpo_rtf(texto_memorial_completo)
                    
                    if corpo_bloco:
                        # Adiciona quebra de página RTF (\page), quebra de parágrafo (\par) e o corpo do bloco
                        # \pard reseta a formatação do parágrafo anterior
                        texto_final_corpo += r'\pard\page\par ' + corpo_bloco
                    else:
                        print(f"AVISO: Corpo RTF vazio ou não detectado no Bloco {idx + 1}. Pulando a anexação.")

                except Exception as e:
                    print(f"ERRO: Falha ao anexar o Bloco {idx + 1} (RTF). Detalhes: {e}. O processo continuará.")
                
        # 4. Monta o arquivo RTF final e salva.
        texto_final_envelopado = texto_final_corpo + "\n" + rodape_rtf
        
        with open(saida_caminho, 'w', encoding='windows-1252', errors='replace') as f:
            f.write(texto_final_envelopado)
            
        print(f"DEBUG: Arquivo final RTF salvo em: {saida_caminho}")
        return

    # --- Lógica DOCX (Criação do Documento Base e Anexação de Conteúdo) ---
    elif extensao == '.docx':
        if Document is MockDocument:
            raise Exception("A biblioteca 'python-docx' não foi encontrada. Instale-a para processar arquivos DOCX (pip install python-docx).")
        
        print(f"\n--- Iniciando Geração do DOCX de Múltiplos Memoriais ({len(todos_conjuntos)} total) ---")
        
        # 1. Processa o PRIMEIRO conjunto no documento base
        dados_gerais_0, pontos_0 = todos_conjuntos[0]
        try:
            documento_final = processar_docx_bloco(caminho_modelo, dados_gerais_0, pontos_0, ignorar_confrontante_repetido)
        except Exception as e:
            raise Exception(f"Falha fatal ao processar o Bloco 1 (DOCX): {e}")
        
        # 2. Processa os conjuntos subsequentes e anexa o conteúdo
        if len(todos_conjuntos) > 1:
            for idx in range(1, len(todos_conjuntos)):
                
                item_conjunto = todos_conjuntos[idx]
                
                if not (isinstance(item_conjunto, (list, tuple)) and len(item_conjunto) == 2):
                    print(f"ERRO GRAVE DE DADOS (BLOCO {idx + 1}): Estrutura corrompida. O processo continuará.")
                    continue
                
                dados_gerais, pontos = item_conjunto
                
                if not dados_gerais and not pontos:
                    print(f"AVISO: Bloco {idx + 1} vazio. Pulando.")
                    continue

                try:
                    print(f"DEBUG: Anexando Bloco {idx + 1} no DOCX.")
                    
                    # Adiciona uma quebra de página
                    documento_final.add_page_break()
                    
                    # Processa o bloco em um documento temporário para reter a formatação
                    documento_bloco = processar_docx_bloco(caminho_modelo, dados_gerais, pontos, ignorar_confrontante_repetido)
                    
                    # CÓPIA SEGURA DO CONTEÚDO
                    elementos_do_bloco = list(documento_bloco.element.body)

                    # Tenta evitar o último elemento, que costuma ser sectionProperties e pode causar erro
                    if len(elementos_do_bloco) > 1 and elementos_do_bloco[-1].tag.endswith('sectPr'):
                        itens_para_copiar = elementos_do_bloco[:-1]
                    else:
                        itens_para_copiar = elementos_do_bloco

                    for elemento in itens_para_copiar:
                        # deepcopy do elemento _element (lxml element) e anexa ao body do documento final
                        documento_final.element.body.append(deepcopy(elemento))

                except Exception as e:
                    # Este erro é a falha na lógica de anexação/processamento DOCX
                    print(f"ERRO: Falha ao anexar o Bloco {idx + 1} (DOCX). Detalhes: {e}. O processo continuará.")
                    
        # 3. Salva o documento DOCX final
        documento_final.save(saida_caminho)
        print(f"DEBUG: Arquivo final DOCX salvo em: {saida_caminho}")
        return

    else:
        raise ValueError(f"Extensão de arquivo não suportada: {extensao}. Suportadas: .rtf, .docx")


# --- Funções de Interface (GUI e CLI) ---
def selecionar_arquivos_e_processar():
    root = tk.Tk()
    root.withdraw()

    try:
        file_types = [
            ("Modelos de Documento", "*.rtf *.docx"),
            ("RTF Files", "*.rtf"),
            ("DOCX Files", "*.docx")
        ]
        
        modelo_caminho = filedialog.askopenfilename(title="1. Selecione o Modelo (RTF ou DOCX)", filetypes=file_types)
        if not modelo_caminho:
            messagebox.showinfo("Aviso", "Seleção cancelada.")
            return
        
        extensao_modelo = Path(modelo_caminho).suffix.lower()
        if extensao_modelo not in ('.rtf', '.docx'):
            raise ValueError("Extensão de modelo não reconhecida. Por favor, use .rtf ou .docx.")
        
        output_file_types = [
            ("Arquivo Final", f"*{extensao_modelo}"),
            ("RTF Files", "*.rtf"),
            ("DOCX Files", "*.docx")
        ]
        
        csv_unico = filedialog.askopenfilename(title="2. Selecione o CSV ÚNICO (Múltiplos Conjuntos com **)", filetypes=[("CSV Files", "*.csv")])
        if not csv_unico:
            messagebox.showinfo("Aviso", "Seleção cancelada.")
            return

        saida_caminho = filedialog.asksaveasfilename(
            title="3. Salvar arquivo final (Múltiplos Memoriais em 1 Arquivo)", 
            defaultextension=extensao_modelo, 
            filetypes=output_file_types
        )
        if not saida_caminho:
            messagebox.showinfo("Aviso", "Seleção cancelada.")
            return

        print("--- Iniciando carregamento do arquivo CSV ÚNICO ---")
        # Chamada da função unificada
        todos_conjuntos = carregar_todos_os_conjuntos(csv_unico)
        
        print(f"\n--- Iniciando o processamento do Modelo ({extensao_modelo}) para {len(todos_conjuntos)} memorial(is) ---")
        
        # Chamada da função orquestradora (padrão para GUI é NÃO ignorar a repetição)
        processar_modelo(modelo_caminho, todos_conjuntos, saida_caminho, ignorar_confrontante_repetido=False)
        
        messagebox.showinfo("Sucesso", f"{len(todos_conjuntos)} Memorial(is) gerado(s) com sucesso em um único arquivo:\n{saida_caminho}")
    except Exception as e:
        # Apenas mostra o erro na tela (GUI)
        print(f"ERRO: {e}")
        messagebox.showerror("Erro", f"Erro ao gerar arquivo:\n{e}")

def main():
    parser = argparse.ArgumentParser(description="Gerador de Memorial Descritivo DOCX/RTF (Modo Múltiplos CSV em 1 Saída)")
    parser.add_argument("--x", action="store_true", 
                        help="Ativa a substituição/remoção de confrontantes repetidos.")
                        
    parser.add_argument("modelo", nargs="?", help="Caminho do Modelo (RTF ou DOCX)")
    parser.add_argument("csv_unico_ou_base", nargs="?", metavar="CSV_BASE", 
                        help="Caminho base do CSV (ex: D:\\TESTE) ou caminho completo (ex: D:\\TESTE-MOD.csv).")
    parser.add_argument("--saida", help="Caminho para salvar o arquivo final")
    
    args = parser.parse_args()
    
    ignorar_rep = args.x
    
    if args.modelo and Path(args.modelo).suffix.lower() == '.docx' and Document is MockDocument:
             print("AVISO: 'python-docx' não está instalado. O processamento de arquivos DOCX NÃO funcionará.", file=sys.stderr)

    try:
        # --- MODO SIMPLIFICADO (Modelo e CSV_BASE) ---
        if args.modelo and args.csv_unico_ou_base and not (args.saida):
            modelo_path = Path(args.modelo)
            extensao = modelo_path.suffix.lower() 
            if not extensao: extensao = ".rtf" 
            
            csv_base_path = Path(args.csv_unico_ou_base) 

            # LÓGICA DE CORREÇÃO DO NOME DO CSV
            if csv_base_path.suffix.lower() != ".csv":
                 # Se o usuário passou 'TESTE', o arquivo de entrada é 'TESTE-MOD.csv'
                 csv_unico = str(csv_base_path.with_name(csv_base_path.name + "-MOD.csv"))
                 # A saída deve manter a base original, ou seja, 'TESTE.docx'
                 saida_base_path = csv_base_path
            else:
                # Se o usuário passou 'TESTE-MOD.csv'
                csv_unico = str(csv_base_path)
                # A saída será 'TESTE-MOD.docx'
                saida_base_path = csv_base_path.with_suffix('') # Remove a extensão .csv

            saida_caminho = str(saida_base_path.with_suffix(extensao)) 

            print(f"--- MODO SIMPLIFICADO DETECTADO ({'Remover Repetição: SIM' if ignorar_rep else 'Remover Repetição: NÃO'}) ---")
            print(f"Modelo : {args.modelo}")
            print(f"CSV de Entrada: {csv_unico}")
            print(f"Saída Esperada: {saida_caminho}")

            todos_conjuntos = carregar_todos_os_conjuntos(csv_unico)
            
            processar_modelo(args.modelo, todos_conjuntos, saida_caminho, ignorar_confrontante_repetido=ignorar_rep)
            print(f"\n✅ {len(todos_conjuntos)} Memorial(is) gerado(s) com sucesso em: {saida_caminho}")

        # --- MODO COMPLETO (Flags explícitas) ---
        elif args.modelo and args.csv_unico_ou_base and args.saida: 
            print(f"--- MODO COMPLETO DETECTADO ({'Remover Repetição: SIM' if ignorar_rep else 'Remover Repetição: NÃO'}) ---")
            
            # Aqui, assumimos que o usuário passou o caminho correto do CSV para o argumento posicional
            csv_path = Path(args.csv_unico_ou_base)
            if csv_path.suffix.lower() != ".csv":
                # Se não tem CSV, tenta completar com -MOD.csv (comportamento de fallback)
                csv_unico = str(csv_path.with_name(csv_path.name + "-MOD.csv"))
                print(f"AVISO: O CSV não tinha extensão. Assumindo: {csv_unico}")
            else:
                csv_unico = str(csv_path)

            todos_conjuntos = carregar_todos_os_conjuntos(csv_unico)

            processar_modelo(args.modelo, todos_conjuntos, args.saida, ignorar_confrontante_repetido=ignorar_rep)
            print(f"\n✅ {len(todos_conjuntos)} Memorial(is) gerado(s) com sucesso em: {args.saida}")

        # --- MODO GUI (fallback) ---
        else:
            print("--- MODO GRÁFICO (GUI) DETECTADO ---")
            selecionar_arquivos_e_processar()

    except Exception as e:
        print(f"ERRO FATAL: Ocorreu um erro durante o processamento. Detalhes: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()